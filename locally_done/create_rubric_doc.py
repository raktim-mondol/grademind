from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_rubric():
    document = Document()

    # Title
    title = document.add_heading('COMP9414 Assignment 2 - Rubric', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('Neural Networks, Tree-based and Ensemble Methods')
    
    # Introduction
    document.add_heading('Overview', level=1)
    document.add_paragraph(
        'This rubric outlines the marking criteria for Assignment 2. '
        'Total Marks: 25'
    )

    # Task 1
    document.add_heading('Task 1: Data Preprocessing (2 Marks)', level=1)
    
    table1 = document.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Sub-task'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Marks'

    data1 = [
        ('1.1', 'Missing data removal or imputation', '1.0'),
        ('1.2', 'Feature and class encoding', '0.5'),
        ('1.3', 'Rescaling attributes', '0.5')
    ]
    for sub, desc, mark in data1:
        row_cells = table1.add_row().cells
        row_cells[0].text = sub
        row_cells[1].text = desc
        row_cells[2].text = mark

    # Task 2
    document.add_heading('Task 2: Model Training (3 Marks)', level=1)
    
    table2 = document.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Sub-task'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Marks'

    data2 = [
        ('2.1', 'Shallow neural net for classification', '1.0'),
        ('2.2', 'Tree-based and Ensemble Models for Classification', '2.0')
    ]
    for sub, desc, mark in data2:
        row_cells = table2.add_row().cells
        row_cells[0].text = sub
        row_cells[1].text = desc
        row_cells[2].text = mark

    # Task 3
    document.add_heading('Task 3: Report (20 Marks)', level=1)
    
    table3 = document.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'Sub-task'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Marks'

    data3 = [
        # 3.1
        ('3.1', 'Preprocessing Analysis:\n- Explain why rescaling is not necessary for tree-based models.\n- Explain why rescaling is necessary for neural networks.\n- Explain what happens if we do not rescale in neural networks.', '2.0'),
        
        # 3.2
        ('3.2.1', 'Model Fine Tuning - Plots:\nCreate separate plots for each model (DT, RF, Bagging, AdaBoost) showing accuracy and loss changes on the test set as hyperparameters vary.', '2.0'),
        ('3.2.2', 'Model Fine Tuning - Analysis:\nExplain how training and test accuracies change as model complexity increases.', '1.0'),
        ('3.2.3', 'Model Fine Tuning - Overfitting:\nDiscuss observations regarding overfitting and explain reasoning.', '1.0'),
        
        # 3.3
        ('3.3.1', 'Imbalanced Data - Detection Code:\nDevelop code to determine if a dataset is imbalanced.', '1.0'),
        ('3.3.2', 'Imbalanced Data - Identification:\nTest code on 3 datasets and identify the imbalanced one.', '1.0'),
        ('3.3.3', 'Imbalanced Data - Performance Table:\nReport performance (confusion matrix, accuracy, precision, recall, F1) of models on the imbalanced dataset.', '2.0'),
        ('3.3.4', 'Imbalanced Data - Handling:\nApply class weighting or resampling, test improvement, and report results in a plot comparing performance.', '2.0'),
        ('3.3.5', 'Imbalanced Data - Discussion:\nDiscuss which model handles class imbalance better and explain why.', '1.0'),

        # 3.4
        ('3.4.1', 'Noisy Labels - Code:\nDevelop code to randomly flip n% (20%) of training labels in adult income dataset.', '1.0'),
        ('3.4.2', 'Noisy Labels - Retraining & Comparison:\nRetrain models on noisy data. Compare performance (metrics + plots) vs original models.', '2.0'),
        ('3.4.3', 'Noisy Labels - Robustness:\nIdentify which model(s) is more robust to noise and explain why.', '1.0'),

        # 3.5
        ('3.5.1', 'NN Fine Tuning - Experiments:\nPlot training/test accuracies/losses while modifying number of neurons. Determine best setting based on accuracy.', '1.0'),
        ('3.5.2', 'NN Fine Tuning - Overfitting:\nObserve and explain overfitting signs in the plots.', '1.0'),
        ('3.5.3', 'NN Fine Tuning - Optimal Neurons:\nDetermine optimal number of neurons and discuss consistency with rule of thumb.', '1.0'),
    ]

    for sub, desc, mark in data3:
        row_cells = table3.add_row().cells
        row_cells[0].text = sub
        row_cells[1].text = desc
        row_cells[2].text = mark

    # Save
    document.save('Assignment_2_Rubric.docx')
    print("Rubric generated successfully: Assignment_2_Rubric.docx")

if __name__ == '__main__':
    create_rubric()
